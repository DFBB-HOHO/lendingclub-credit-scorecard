#!/usr/bin/env python3
"""Build a polished Chinese model-development report from generated artefacts."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTIFACTS = ROOT / "artifacts"
FIGURES = ROOT / "reports" / "figures"
OUTPUT = ROOT / "reports" / "郭锦航_LendingClub贷前信用评分卡项目报告.docx"

NAVY = "18324A"
TEAL = "2A7F8E"
GOLD = "C89B3C"
LIGHT = "F2F4F7"
PALE_TEAL = "E8F2F3"
RED = "A6453D"
GREY = "606A73"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360


def set_run_font(run, size=10.5, bold=False, color="202428", italic=False, latin="Noto Sans CJK SC", east_asia="Noto Sans CJK SC"):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=GREY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=GREY)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Body Text Custom")
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        r2 = p.add_run(text[len(bold_lead) :])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="Bullet Custom")
    r = p.add_run("•  " + text)
    set_run_font(r, size=10.3)
    return p


def add_callout(doc, title, text, fill=PALE_TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_dataframe_table(doc, frame, widths, header_fill=LIGHT, font_size=9.2, formats=None):
    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, column in enumerate(frame.columns):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(column))
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    formats = formats or {}
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for idx, (column, value) in enumerate(row.items()):
            formatter = formats.get(column)
            shown = formatter(value) if formatter else ("—" if pd.isna(value) else str(value))
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(shown)
            set_run_font(r, size=font_size)
    return table


def add_picture(doc, path, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, italic=True, color=GREY)


def set_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(10.5)

    body = doc.styles.add_style("Body Text Custom", WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = normal
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.10
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    bullet = doc.styles.add_style("Bullet Custom", WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = normal
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(5)
    bullet.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, TEAL, 16, 8),
        ("Heading 2", 13, NAVY, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    metrics = json.loads((ARTIFACTS / "metrics.json").read_text(encoding="utf-8"))
    splits = pd.read_csv(ARTIFACTS / "sample_split.csv")
    iv = pd.read_csv(ARTIFACTS / "iv_summary.csv")
    feature_psi = pd.read_csv(ARTIFACTS / "feature_psi.csv")
    strategy = pd.read_csv(ARTIFACTS / "approval_strategy_scenarios.csv")
    lift = pd.read_csv(ARTIFACTS / "oot_decile_lift.csv")
    decisions = pd.read_csv(ARTIFACTS / "feature_decisions.csv")
    config = json.loads((ROOT / "configs" / "model_config.json").read_text(encoding="utf-8"))
    selected = decisions.loc[decisions["decision"] == "keep", "feature"].tolist()
    oot = metrics["oot"]
    val = metrics["validation"]
    score_psi = metrics["stability"]["development_to_oot_score_psi"]

    doc = Document()
    set_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("LENDINGCLUB 贷前信用评分卡  |  模型开发与风险策略报告")
    set_run_font(r, size=8.5, bold=True, color=GREY)
    add_page_number(section.footer.paragraphs[0])

    # Editorial cover.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PERSONAL CREDIT RISK PROJECT")
    set_run_font(r, size=10, bold=True, color=GOLD)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LendingClub 贷前申请信用评分卡")
    set_run_font(r, size=28, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("从数据治理、WOE 分箱到 OOT 验证、PD 校准与策略阈值")
    set_run_font(r, size=14, color=TEAL)
    p.paragraph_format.space_after = Pt(26)

    strip = doc.add_table(rows=2, cols=3)
    strip.style = "Table Grid"
    set_table_geometry(strip, [3120, 3120, 3120])
    set_repeat_table_header(strip.rows[0])
    labels = ["OOT AUC", "OOT KS", "SCORE PSI"]
    values = [f"{oot['auc']:.3f}", f"{oot['ks']:.1%}", f"{score_psi:.3f}"]
    for idx in range(3):
        shade_cell(strip.cell(0, idx), NAVY)
        shade_cell(strip.cell(1, idx), PALE_TEAL)
        strip.cell(0, idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        strip.cell(1, idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = strip.cell(0, idx).paragraphs[0].add_run(labels[idx])
        set_run_font(r, size=9, bold=True, color=WHITE)
        r = strip.cell(1, idx).paragraphs[0].add_run(values[idx])
        set_run_font(r, size=18, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(68)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("郭锦航  |  金融科技")
    set_run_font(r, size=11, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("面向风险策略分析 / 量化建模岗位  |  2026 年 7 月")
    set_run_font(r, size=9.5, color=GREY)
    doc.add_page_break()

    add_heading(doc, "结论先行：方向合理，但必须从“课程作业”升级成风险项目", 1)
    add_callout(
        doc,
        "判断",
        f"信用评分卡与你现有的量化实习经历形成有效互补，值得写进简历；但只完成“清洗—WOE—逻辑回归—KS/AUC”仍过于标准化。升级版加入时间窗治理、泄露审查、OOT 验证、截距校准、PSI 监控与策略阈值，才足以支撑微众银行风险策略/量化建模岗的面试追问。",
    )
    assessment = pd.DataFrame(
        [
            ["样本切分", "随机切分", "2014–2015 开发 / 2016 验证 / 2017 OOT"],
            ["变量治理", "按缺失率删列", "仅申请时点变量；剔除 grade、利率及贷后字段"],
            ["评估", "AUC、KS", "AUC、KS、十分位 Lift、PD 校准、Score/Feature PSI"],
            ["策略", "给一个 cutoff", "输出不同目标通过率下的 cutoff 与观察坏账率"],
            ["边界", "默认可用于审批", "明确 accepted-only 样本与 reject inference 局限"],
        ],
        columns=["维度", "基础版", "本项目升级版"],
    )
    add_dataframe_table(doc, assessment, [1500, 2700, 5160], font_size=9.2)

    add_heading(doc, "1. 数据与样本设计", 1)
    add_body(
        doc,
        "数据来自 LendingClub 2007–2018Q4 公开借贷记录。原始文件约 226 万条，保留 Fully Paid、Charged Off、Default 三类已形成最终结果的贷款，定义 Default=1（Charged Off / Default）与 Default=0（Fully Paid）。Current、Late、In Grace Period 不被强行贴标签。",
    )
    add_body(
        doc,
        "开发窗口采用 2014–2015。更早的 2007–2013 仍用于 vintage 诊断，但由于平台政策与细粒度征信字段覆盖口径不可比，不参与拟合；2018 年观察坏账率随月份异常下降，判定存在明显右删失，仅保留在诊断中。",
    )
    split_display = splits.copy()
    split_display["sample"] = split_display["sample"].map(
        {
            "train": "开发",
            "validation": "验证",
            "oot": "OOT",
            "excluded_legacy_vintages": "历史诊断",
            "excluded_right_censored": "右删失诊断",
        }
    )
    split_display["bad_rate"] = split_display["bad_rate"].map(lambda x: f"{x:.2%}")
    split_display["observations"] = split_display["observations"].map(lambda x: f"{int(x):,}")
    split_display["defaults"] = split_display["defaults"].map(lambda x: f"{int(x):,}")
    split_display.columns = ["样本", "起始", "结束", "样本量", "违约数", "坏账率"]
    add_dataframe_table(doc, split_display, [2200, 1500, 1500, 1500, 1300, 1360], font_size=8.6)
    add_picture(doc, FIGURES / "vintage_diagnostics.png", "图 1  Vintage 数量与已观察坏账率：2018 年右删失信号清晰")

    add_heading(doc, "2. 泄露审查与特征工程", 1)
    add_bullet(doc, "严格排除 grade、sub_grade、int_rate：它们是 LendingClub 既有授信判断的输出，放进新模型会形成 target leakage / policy leakage。")
    add_bullet(doc, "排除 funded_amnt、installment、recoveries、last_pymnt_*、total_pymnt_* 等放款或贷后字段。")
    add_bullet(doc, "排除州与 ZIP Code：即使可能带来 lift，也存在地域代理、公平性与跨市场迁移风险。")
    add_bullet(doc, "构造 loan_to_income 与 credit_history_months；所有缩尾边界只在开发集拟合，再原样应用于验证/OOT。")
    add_bullet(doc, "连续变量采用训练期分位点预分箱与相邻违约率单调合并；类别变量合并低频水平，使用 0.5 平滑计算 WOE/IV。")

    add_heading(doc, "3. 特征筛选与评分卡", 1)
    add_body(doc, f"候选变量共 {len(config['features']['continuous']) + len(config['features']['categorical'])} 个，经 IV≥0.02、IV≤0.50、WOE 相关性≤0.70 与最多 15 个变量的复杂度约束，最终保留 {len(selected)} 个变量。")
    add_body(doc, "最终变量：" + "、".join(selected) + "。")
    iv_display = iv.head(10).copy()
    iv_display["iv"] = iv_display["iv"].map(lambda x: f"{x:.4f}")
    iv_display.columns = ["变量", "IV"]
    add_dataframe_table(doc, iv_display, [6500, 2860], font_size=9.0)
    add_picture(doc, FIGURES / "feature_iv.png", "图 2  开发样本 Information Value 排序")
    add_body(doc, "模型采用 L2 正则 Logistic Regression，以验证集 Log Loss 在 C∈{0.01,0.1,1,10} 中选参；随后只移动截距，使 2016 平均预测 PD 对齐观察坏账率。该校准不改变样本排序、系数相对作用、AUC 或 KS。")
    add_body(doc, "评分映射采用 Base Score=600、Base Odds(Good:Bad)=20:1、PDO=50；分数越高表示风险越低。")

    add_heading(doc, "4. 模型表现", 1)
    perf = pd.DataFrame(
        [
            ["开发 2014–2015", metrics["train"]["auc"], metrics["train"]["ks"], metrics["train"]["bad_rate"], metrics["train"]["mean_predicted_pd"], metrics["train"]["brier_score"]],
            ["验证 2016", val["auc"], val["ks"], val["bad_rate"], val["mean_predicted_pd"], val["brier_score"]],
            ["OOT 2017", oot["auc"], oot["ks"], oot["bad_rate"], oot["mean_predicted_pd"], oot["brier_score"]],
        ],
        columns=["样本", "AUC", "KS", "实际坏账率", "平均预测 PD", "Brier"],
    )
    fmts = {
        "AUC": lambda x: f"{x:.3f}",
        "KS": lambda x: f"{x:.1%}",
        "实际坏账率": lambda x: f"{x:.2%}",
        "平均预测 PD": lambda x: f"{x:.2%}",
        "Brier": lambda x: f"{x:.4f}",
    }
    add_dataframe_table(doc, perf, [2200, 1300, 1300, 1600, 1700, 1260], font_size=8.8, formats=fmts)
    add_picture(doc, FIGURES / "oot_roc_ks.png", "图 3  2017 OOT ROC 曲线与 KS")
    add_picture(doc, FIGURES / "oot_score_distribution.png", "图 4  2017 OOT 好坏客户评分分布")
    high = lift.iloc[0]
    low = lift.iloc[-1]
    add_callout(
        doc,
        "风险排序解释",
        f"OOT 最高风险十分位观察坏账率为 {high['bad_rate']:.1%}，最低风险十分位为 {low['bad_rate']:.1%}，相差约 {high['bad_rate']/low['bad_rate']:.1f} 倍；校准后 OOT 平均预测 PD 为 {oot['mean_predicted_pd']:.2%}，与实际 {oot['bad_rate']:.2%} 接近。",
    )

    add_heading(doc, "5. 稳定性与监控", 1)
    add_body(doc, f"开发到 OOT 的 Score PSI={score_psi:.3f}，按常用经验阈值属于稳定区间。总分稳定不代表所有底层字段稳定，因此同时输出单变量 PSI；生产环境应将字段覆盖率、缺失率与 PSI 联合监控。")
    psi_display = feature_psi.head(10).copy()
    psi_display["psi"] = psi_display["psi"].map(lambda x: f"{x:.4f}")
    psi_display.columns = ["变量", "PSI", "状态"]
    add_dataframe_table(doc, psi_display, [5000, 2000, 2360], font_size=9.0)

    add_heading(doc, "6. 策略阈值情景（非真实拒绝策略）", 1)
    add_body(doc, "下表在历史已放款客户内部按分数由高到低截取，用于展示风险—规模 trade-off。由于没有拒绝客户的真实违约结果，这些“通过率”只能称为情景分析，不可直接解释为新客群的真实审批效果。")
    strategy_pick = strategy.iloc[[2, 4, 6, 8]].copy()
    strategy_pick = strategy_pick[["target_approval_rate", "score_cutoff", "approved_bad_rate", "default_capture_in_approved"]]
    strategy_pick.columns = ["目标通过率", "分数阈值", "通过客群坏账率", "通过客群包含的违约占比"]
    fmts = {
        "目标通过率": lambda x: f"{x:.0%}",
        "分数阈值": lambda x: f"{x:.1f}",
        "通过客群坏账率": lambda x: f"{x:.2%}",
        "通过客群包含的违约占比": lambda x: f"{x:.2%}",
    }
    add_dataframe_table(doc, strategy_pick, [1800, 1800, 2600, 3160], font_size=9.0, formats=fmts)

    add_heading(doc, "7. 模型边界与生产化缺口", 1)
    for item in (
        "样本只有历史已放款客户，存在 selection bias；Rejected Loans 没有真实表现标签，不能声称已完成 reject inference。",
        "美国 P2P 个人贷样本不能直接迁移为中国数字银行模型，项目价值在于可复现的方法、治理意识和策略翻译能力。",
        "公开数据缺少欺诈标签、设备与行为序列、收入核验、LGD/EAD、资金成本和运营成本，不能据此计算真实利润最大化阈值。",
        "正式投产还需要公平性测试、模型审批、版本管理、champion/challenger、监控阈值及回滚机制。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "8. 可直接用于简历的项目表述", 1)
    add_callout(doc, "项目名称", "LendingClub 贷前申请评分卡与风险策略监控  |  Python / WOE / Logistic Regression  |  2026.07")
    resume_bullets = [
        "样本与治理：对 226 万条公开借贷记录进行状态、时间窗与字段可用性审计，形成约 134.5 万条已结清/违约表现样本；识别 2018 年右删失，采用 2014–2015 开发、2016 验证、2017 OOT，避免随机切分造成时间穿越。",
        f"建模与评分：对 {len(config['features']['continuous']) + len(config['features']['categorical'])} 个贷前申请/征信变量完成训练期缩尾、单调 WOE 分箱、IV 与相关性筛选，保留 15 项变量构建 Logistic 评分卡，并完成 600 分、20:1 Base Odds、PDO50 的分数映射与验证期截距校准。",
        f"效果与策略：2017 OOT AUC={oot['auc']:.3f}、KS={oot['ks']:.1%}、Score PSI={score_psi:.3f}，最高/最低风险十分位坏账率为 {high['bad_rate']:.1%}/{low['bad_rate']:.1%}；输出 PSI 监控、Lift 与不同目标通过率下的 cutoff 情景，并明确 accepted-only 样本偏差。",
    ]
    for item in resume_bullets:
        add_bullet(doc, item)

    add_heading(doc, "9. 面试追问口径", 1)
    qa = [
        ("为什么不用随机切分？", "信贷风险随宏观环境、平台政策和客群结构漂移。随机切分会把未来分布泄露进训练，使指标虚高；OOT 更接近真实上线。"),
        ("为什么删掉 grade 和利率？", "它们是 LendingClub 原有风控决策的结果，等于把旧模型答案喂给新模型，会形成 policy leakage。"),
        ("为什么选 Logistic 而不是只做 XGBoost？", "评分卡需要稳定、可解释、易审计和可映射为分数。树模型可以作为 challenger，但不应替代基线治理。"),
        ("AUC 0.67 是否太低？", "项目严格限制为申请时点变量并做 OOT；指标是诚实的跨期表现。风险模型还要看稳定性、校准、可解释性和策略价值，不能只追单一 AUC。"),
        ("PD 为什么要校准？", "排序正确不等于总体风险水平正确。2016 客群坏账率上升，只移动截距可修正先验坏账率，同时保留所有排序和变量相对作用。"),
        ("为什么不能直接用 rejected 数据？", "Rejected Loans 没有真实违约标签；把“被拒”当成“违约”是伪标签。需要 reject inference 假设或后续可观测表现，且必须披露偏差。"),
        ("策略阈值怎么定？", "本项目只展示通过率—坏账率情景。生产中还要把 PD 与 LGD、EAD、利率、资金与运营成本相连，按预期利润和风险偏好定阈值。"),
        ("上线后监控什么？", "Score/Feature PSI、缺失率、字段覆盖率、AUC/KS、校准、Vintage、审批率与坏账率；超过阈值触发排查、重校准或重训。"),
    ]
    for question, answer in qa:
        p = doc.add_paragraph(style="Body Text Custom")
        r = p.add_run(question + " ")
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(answer)
        set_run_font(r)

    add_heading(doc, "10. 复现与交付结构", 1)
    add_body(doc, "一键复现：先运行 `python scripts/download_data.py` 下载并校验数据，再运行 `python scripts/train_scorecard.py` 生成模型、分箱、指标、PSI、策略表与图形。单元测试命令为 `PYTHONPATH=src python -m unittest discover -s tests -v`。")
    for item in (
        "src/credit_scorecard：数据治理、WOE 分箱、特征筛选、Logistic、评分映射、监控与报告模块",
        "configs/model_config.json：数据、时间窗、特征、分箱、筛选、模型和评分参数",
        "artifacts：分箱规则、IV/系数、scorecard points、AUC/KS、PSI、Lift、策略阈值",
        "reports：Markdown 模型报告、图形与本 Word 报告",
        "tests：分箱单调性、未见类别、评分可逆性、PSI、截距校准测试",
    ):
        add_bullet(doc, item)

    add_heading(doc, "数据来源", 1)
    sources = [
        "Kaggle：All Lending Club loan data — https://www.kaggle.com/datasets/wordsforthewise/lending-club",
        "DePaul University 公共镜像与数据字典 — https://bigblue.depaul.edu/jlee141/econdata/LendingClub_LoanData/",
        "UCM / Zenodo 贷前授信治理版数据说明（用于字段泄露交叉校验）— https://doi.org/10.5281/zenodo.11295916",
    ]
    for source in sources:
        add_bullet(doc, source)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
